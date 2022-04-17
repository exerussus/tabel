from function import json_read_text_file


def print_tabel():
    """Печатает табель рабочих часов"""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    data_replace = json_read_text_file('data_replace')
    data_teachers = json_read_text_file('data_teachers')
    month = input('Введите номер месяца: ')
    year = '2022'
    doc = docx.Document()
    title = doc.add_paragraph('') # Создание абзаца
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
    title = title.add_run('Расшифровка к табелю на заработную плату\n'  # Добавление текста в абзац
                          f'за {month} {year} г. (первичная) по МБОУ «Школа № 38»  г. \n'
                          'Рязани\n'
                          'Оплатить:\n')
    title.bold = True  # Изменение стиля текста на жирный
    title.font.name = 'Times New Roman'  # Изменение стиля текста на Times New Roman
    title.font.size = Pt(14)  # Изменение размера текста

    title2 = doc.add_paragraph('')
    title2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title2 = title2.add_run(data_teachers['моисеева']['дательный'])
    title2.font.name = 'Times New Roman'  # Изменение стиля текста на Times New Roman
    title2.font.size = Pt(12)  # Изменение размера текста

    data_action = {}
    for day in data_replace[month]:
        for for_teacher, replace_values in day.items():
            data_action[replace_values[2]] = [replace_values[2], [for_teacher], replace_values[0], replace_values[1]]

    # добавляем таблицу
    table = doc.add_table(rows=2, cols=7)  # rows - строки, cols - столбцы
    # применяем стиль для таблицы
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = '1'
    cell2 = table.cell(0, 1)
    cell2.text = 'За Ярцеву И.И.'
    cell3 = table.cell(0, 2)
    cell3.text = 'класс'
    cell4 = table.cell(0, 3)
    cell4.text = 'предмет'
    cell5 = table.cell(0, 4)
    cell5.text = 'дата'
    cell7 = table.cell(0, 6)
    cell7.text = 'часы'
    # # заполняем таблицу данными
    # for row in range(1):
    #     x = 0
    #     for col in range(7):
    #         # получаем ячейку таблицы
    #         cell = table.cell(row, col)
    #         # записываем в ячейку данные
    #         cell.text = str(row + data_replace[month]) + str(col + 1)

    doc.save('example.docx')


