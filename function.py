


#
# def days_generation_for_month(data_base, month, debug=None):
#     """Заполняет месяц днями. Всего будет 31 дней независимо от месяца.
#     Аргументы: data_base - заполняемый словарь; month - заполняемый месяц в словаре."""
#     if debug:
#         print('На входе:')
#         print(data_base)
#     str_month = str(month)
#     if len(str_month) == 1:
#         str_month = '0' + str_month
#     for i in range(31):
#         i = str(i + 1)
#         if len(i) == 1:
#             i = '0' + i
#         data_base[month][f'{i}.' + str_month] = ''
#     if debug:
#         print('На выходе:')
#         print(data_base)


def do_title(inputting_date):
    """Создаёт титл на день
    Аргументы: inputting_date - введённая пользователем дата (день|месяц)"""
    import datetime as dt
    from data import days
    now = dt.datetime.now()
    year_now = now.year
    m = inputting_date.split('.')
    try:
        day = m[0]
        month = m[1]
    except IndexError:
        month = m[0][3:]
        day = m[0][:2]
    day = int(day)
    month = int(month)
    title_day = dt.datetime(year_now, month, day)
    day_of_week = days['родительный'][title_day.strftime('%w')]
    return f'Замена на {day_of_week} ({inputting_date})'


class Timetable:

    def __init__(self, order):
        """Инициация строки типа:
            '31.01, за ярцева, моисеева, 6в, 3ур, 1ч'"""

        self.order = order

    def decypher(self):
        """Разделяет текст типа "31.01, за ярцева, моисеева, 6в, 3ур, 1ч" на отдельные элементы:
            'date': date,
            'replacing_teacher' - учитель которого заменяют,
            'teacher' - заменяющий учитель,
            'index_class' - номер и буква класса,
            'index_subject' - номер урока,
            'hours' - часы,
        """
        x = self.order.split(', ')
        date = x[0]
        replacing_teacher = x[1][3:]
        teacher = x[2]
        index_class = x[3]
        index_subject = x[4][:1]
        hours = x[5][:1]
        y = {
            'date': date,
            'replacing_teacher': replacing_teacher,
            'teacher': teacher,
            'index_class': index_class,
            'index_subject': index_subject,
            'hours': hours,
        }
        return y

    def for_teacher(self):
        """За кого замена"""
        from data import data_teachers
        teacher = self.decypher()['replacing_teacher']
        x = 'За ' + data_teachers[teacher]['родительный'] + ':'
        return x

    def making_replace(self):
        """Заполняет словарь с заменами"""
        data_teachers = JsonOperations.read('data_teachers')
        data_replace = JsonOperations.read('data_replace')
        x = self.decypher()
        date = x['date']
        split_date = date.split('.')
        month = split_date[1]
        index_subject = x['index_subject']
        index_class = x['index_class']
        teacher = x['teacher']
        join = [index_subject, index_class, data_teachers[teacher]['именительный']]
        """[index_subject, index_class, data_teachers[teacher]['именительный']]"""
        for_teacher = self.for_teacher()

        try:
            data_replace[month][date][for_teacher][index_subject] = join
        except KeyError:
            try:
                data_replace[month][date][for_teacher] = {index_subject: join}
            except KeyError:
                days_generation_for_month(data_replace, month)
                data_replace[month][date] = {for_teacher: {'1': '', '2': '', '3': '', '4': '', '5': '', '6': ''}}
                data_replace[month][date][for_teacher][index_subject] = join
        JsonOperations.save('data_replace', data_replace)
        JsonOperations.save('data_teachers', data_teachers)


def print_timetable():
    """Выводит таблицу с заменами"""
    data_replace = JsonOperations.read('data_replace')
    date = input('Введите дату: ')

    try:
        month = date[3:]
        day = date[:2]
        int(day)
        print(do_title(date))
        try:
            for a, b in data_replace[month][date].items():
                print(a)
                for i in b:
                    if b[i] != '':
                        x = b[i][0] + ') ' + b[i][1] + ' ' + b[i][2]
                        print(x)
                print('')
        except KeyError:
            print('Нет изменений на данную дату.')
    except ValueError:
        print('Неверно введена дата...')


def input_replace():
    """Вводит данные в data_replace"""
    while True:
        print('Для выхода нажмите 0.')
        text = input('Введите текст типа "31.01, за ярцева, моисеева, 6в, 3ур, 1ч":')
        if text == '0':
            break
        else:
            try:
                n = Timetable(text)
                n.making_replace()
            except IndexError:
                print("Введена некорректная информация...")


class WordPrinter:

    @staticmethod
    def print_tabel():
        """Печатает табель рабочих часов"""
        import docx
        from docx.shared import Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        data_replace = JsonOperations.read('data_replace')
        data_teachers = JsonOperations.read('data_teachers')
        month = input('Введите номер месяца: ')
        year = '2022'
        doc = docx.Document()
        title = doc.add_paragraph('') # Создание абзаца
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
        title = title.add_run('Расшифровка к табелю на заработную плату\n'  # Добавление текста в абзац
                              f'за {month} {year} г. (первичная) по МБОУ «Школа № 38»  г. \n'
                              'Рязани\n'
                              'Оплатить:\n')
        title.bold = True  # Изменение стиля текста на жирный
        title.font.name = 'Times New Roman'  # Изменение стиля текста на Times New Roman
        title.font.size = Pt(14)  # Изменение размера текста

        title2 = doc.add_paragraph('')
        title2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title2 = title2.add_run(data_teachers['моисеева']['дательный'])
        title2.font.name = 'Times New Roman'  # Изменение стиля текста на Times New Roman
        title2.font.size = Pt(12)  # Изменение размера текста

        # добавляем таблицу
        table = doc.add_table(rows=2, cols=7)  # rows - строки, cols - столбцы
        # применяем стиль для таблицы
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        cell.text = '1'
        cell2 = table.cell(0, 1)
        cell2.text = 'За Ярцеву И.И.'
        cell3 = table.cell(0, 2)
        cell3.text = 'класс'
        cell4 = table.cell(0, 3)
        cell4.text = 'предмет'
        cell5 = table.cell(0, 4)
        cell5.text = 'дата'
        cell7 = table.cell(0, 6)
        cell7.text = 'часы'
        # # заполняем таблицу данными
        # for row in range(1):
        #     x = 0
        #     for col in range(7):
        #         # получаем ячейку таблицы
        #         cell = table.cell(row, col)
        #         # записываем в ячейку данные
        #         cell.text = str(row + data_replace[month]) + str(col + 1)

        doc.save('example.docx')


def main_menu():
    # truth = random.choice(data_truth)
    # print(f'Лизонька, {truth}')
    while True:
        choice = input('   Выберите действие:\n'
                       '1. Добавить изменения\n'
                       '2. Вывести расписание\n'
                       '3. Вывести табель\n'
                       '0. Выход\n'
                       'Ваш выбор: ')

        if choice == '1':
            input_replace()
        elif choice == '2':
            print_timetable()
        elif choice == '3':
            WordPrinter.print_tabel()
        elif choice == '0':
            print('Спасибо, что воспользовались программой Tabel и облегчили себе жизнь!')
            input()
            exit(0)
        else:
            print("Выберите действие из списка...")




